import pandas as pd
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Border, Side, Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image
import matplotlib.pyplot as plt
import io
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import tempfile
import atexit
from datetime import datetime

# Определяем базовую директорию (на уровень выше скрипта)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Получаем текущую дату для имени файла
current_date = datetime.now().strftime("%d.%m.%Y")

# Конфигурация путей к файлам относительно базовой директории
FILE_PATHS = {
    'kr_file': os.path.join(BASE_DIR, "КР", "Проект плана КР 2027_20.xlsx"),
    'totr_file': os.path.join(BASE_DIR, "ТОиТР", "Проект плана ТОиТР 2027.xlsx"),
    'output_file': os.path.join(BASE_DIR, f"Отчет_по_подготовке_ТОиР_2027_{current_date}.docx")
}

# Глобальный список для хранения временных файлов
temp_files = []

def cleanup_temp_files():
    """Очистка временных файлов при завершении программы"""
    for temp_file in temp_files:
        try:
            if os.path.exists(temp_file):
                os.unlink(temp_file)
        except Exception as e:
            print(f"Ошибка при удалении временного файла {temp_file}: {e}")

# Регистрируем функцию очистки при выходе
atexit.register(cleanup_temp_files)

def check_file_exists(file_path, file_description):
    """Проверяет существование файла и выводит информационное сообщение"""
    if not os.path.exists(file_path):
        print(f"Файл не найден: {file_path}")
        print(f"Описание: {file_description}")
        print(f"Текущая рабочая директория: {os.getcwd()}")
        print(f"Базовая директория: {BASE_DIR}")
        print("Пожалуйста, проверьте пути в конфигурации FILE_PATHS")
        return False
    else:
        print(f"Файл найден: {file_path}")
        return True

def save_buffer_to_temp_file(buffer, prefix="chart"):
    """Сохраняет буфер изображения во временный файл и возвращает путь к нему"""
    try:
        # Создаем временный файл
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png', prefix=prefix)
        temp_file.write(buffer.getvalue())
        temp_file.flush()
        temp_files.append(temp_file.name)
        return temp_file.name
    except Exception as e:
        print(f"Ошибка при сохранении временного файла: {e}")
        return None

def generate_kr_report():
    """Генерация отчета по капитальному ремонту"""
    
    # Проверяем существование файла
    if not check_file_exists(FILE_PATHS['kr_file'], "Файл данных по капитальному ремонту"):
        raise FileNotFoundError(f"Файл КР не найден: {FILE_PATHS['kr_file']}")
    
    # Читаем данные
    df = pd.read_excel(
        FILE_PATHS['kr_file'],
        sheet_name="ПроектКР2026",
        usecols="AO,AQ,BD,BJ,BM,BT,BV,CI,CK",
        skiprows=15,
        nrows=1000,
        names=['ПО_Общества', 'План', 'МТР', 'ДВ', 'КП', 'Передано_в_ОДСиССР', 'Направлено_на_осмечивание', 'Статус_объекта', 'Признак_МТР_в_заказе']
    ).dropna(subset=['ПО_Общества'])
    
    # Предварительно обрабатываем значения для правильного отображения
    df['План'] = df['План'].replace({
        'Основная': 'Основной',
        'Доп1': 'Доп_1',
        'Доп2': 'Доп_2'
    })

    df['ДВ'] = df['ДВ'].replace({
        'НА ПРОВЕРКЕ': 'ДВ на проверке',
        'ДА': 'ДВ принята в работу',
        'НЕТ': 'ДВ отсутствует'
    })

    df['КП'] = df['КП'].replace({
        'НА ПРОВЕРКЕ': 'КП на проверке',
        'ДА': 'КП принято в работу',
        'НЕТ': 'КП отсутствует',
        'НЕ требуется': 'КП не требуется'
    })

    df['МТР'] = df['МТР'].replace({
        'НА ПРОВЕРКЕ': 'МТР на проверке',
        'ДА': 'ЕСТЬ замечания к МТР',
        'НЕТ': 'Замечаний к МТР НЕТ',
        'НЕ ТРЕБУЕТСЯ': 'Внесение МТР не требуется'
    })

    df['Признак_МТР_в_заказе'] = df['Признак_МТР_в_заказе'].replace({
        'Да': 'Есть признаки МТР в заказе',
        'Нет': 'Нет признаков МТР в заказе',
        'Не требуется': 'Не требуется МТР в заказе'
    })

    df['Передано_в_ОДСиССР'] = df['Передано_в_ОДСиССР'].replace({
        'ДА': 'Передано в ОДСиССР',
        'НЕТ': 'Не передано в ОДСиССР'
    })

    df['Направлено_на_осмечивание'] = df['Направлено_на_осмечивание'].replace({
        'ДА': 'Направлено на осмечивание',
        'НЕТ': 'Не направлено на осмечивание',
        'На доработке': 'СД на доработке'
    })

    df['Статус_объекта'] = df['Статус_объекта'].replace({
        'НА ПРОВЕРКЕ': 'Объект на проверке',
        'РАЗРАБОТКА СД': 'Разработка СД по объекту',
        'ВКЛ': 'Объект включен в план',
        'ПРЕД. К ИСКЛ': 'Объект предлагается к исключению',
        'ИСКЛ': 'Объект исключен из плана'
    })

    # Создаем сводные таблицы и объединяем
    result = pd.concat([
        pd.crosstab(df['ПО_Общества'], df['План']),
        pd.crosstab(df['ПО_Общества'], df['ДВ']),
        pd.crosstab(df['ПО_Общества'], df['КП']),
        pd.crosstab(df['ПО_Общества'], df['МТР']),
        pd.crosstab(df['ПО_Общества'], df['Признак_МТР_в_заказе']),
        pd.crosstab(df['ПО_Общества'], df['Передано_в_ОДСиССР']),
        pd.crosstab(df['ПО_Общества'], df['Направлено_на_осмечивание']),
        pd.crosstab(df['ПО_Общества'], df['Статус_объекта'])
    ], axis=1, sort=False)    

    # Добавляем настоящие столбцы
    required_columns = ['Основной', 'Доп_1', 'Доп_2',
                       'ДВ на проверке', 'ДВ принята в работу', 'ДВ отсутствует',
                       'КП на проверке', 'КП принято в работу', 'КП отсутствует', 'КП не требуется',
                       'МТР на проверке', 'ЕСТЬ замечания к МТР', 'Замечаний к МТР НЕТ', 'Внесение МТР не требуется',
                       'Есть признаки МТР в заказе', 'Нет признаков МТР в заказе', 'Не требуется МТР в заказе',
                       'Передано в ОДСиССР', 'Не передано в ОДСиССР',
                       'Направлено на осмечивание', 'Не направлено на осмечивание', 'СД на доработке',
                       'Объект на проверке', 'Разработка СД по объекту', 'Объект включен в план', 'Объект предлагается к исключению', 'Объект исключен из плана']
    for col in required_columns:
        if col not in result.columns:
            result[col] = 0
  
    # Сбрасываем индекс, чтобы 'ПО_Общества' стал обычным столбцом
    result = result.reset_index()

    # Добавляем столбец 'Кол-во объектов' - количество уникальных записей для каждого ПО_Общества
    count_series = df.groupby('ПО_Общества').size().reset_index(name='Кол-во объектов')
    result = result.merge(count_series, on='ПО_Общества', how='left')

    # Формируем итоговую таблицу
    result = result[['ПО_Общества', 'Кол-во объектов'] + required_columns]

    # Добавляем строку с общим итогом
    total_row = result.sum(numeric_only=True)
    total_row['ПО_Общества'] = 'Общий итог'
    result = pd.concat([result, pd.DataFrame([total_row])], ignore_index=True, sort=False)
    
    return result

def generate_totr_report():
    """Генерация отчета по техническому обслуживанию и текущему ремонту"""
    
    # Проверяем существование файла
    if not check_file_exists(FILE_PATHS['totr_file'], "Файл данных по техническому обслуживанию и текущему ремонту"):
        raise FileNotFoundError(f"Файл ТОиТР не найден: {FILE_PATHS['totr_file']}")
    
    # Читаем данные
    df = pd.read_excel(
        FILE_PATHS['totr_file'],
        sheet_name="ПроектТОиТР2026",
        usecols="AM,AO,BB,BE,BJ,BL,BU,BW",
        skiprows=15,
        nrows=1500,
        names=['ПО_Общества', 'План', 'ДВ', 'КП', 'Передано_в_ОДСиССР', 'Направлено_на_осмечивание', 'Статус_объекта', 'Признак_МТР_в_заказе']
    ).dropna(subset=['ПО_Общества'])
    
    # Предварительно обрабатываем значения для правильного отображения
    df['План'] = df['План'].replace({
        'Основная': 'Основной',
        'Доп1': 'Доп_1',
        'Доп2': 'Доп_2'
    })

    df['ДВ'] = df['ДВ'].replace({
        'НА ПРОВЕРКЕ': 'ДВ на проверке',
        'ДА': 'ДВ принята в работу',
        'НЕТ': 'ДВ отсутствует'
    })

    df['КП'] = df['КП'].replace({
        'НА ПРОВЕРКЕ': 'КП на проверке',
        'ДА': 'КП принято в работу',
        'НЕТ': 'КП отсутствует',
        'Не требуется': 'КП не требуется'
    })

    df['Признак_МТР_в_заказе'] = df['Признак_МТР_в_заказе'].replace({
        'Да': 'Есть признаки МТР в заказе',
        'Нет': 'Нет признаков МТР в заказе',
        'Не требуется': 'Не требуется МТР в заказе'
    })

    df['Передано_в_ОДСиССР'] = df['Передано_в_ОДСиССР'].replace({
        'ДА': 'Передано в ОДСиССР',
        'НЕТ': 'Не передано в ОДСиССР'
    })

    df['Направлено_на_осмечивание'] = df['Направлено_на_осмечивание'].replace({
        'ДА': 'Направлено на осмечивание',
        'НЕТ': 'Не направлено на осмечивание',
        'На доработке': 'СД на доработке'
    })

    df['Статус_объекта'] = df['Статус_объекта'].replace({
        'НА ПРОВЕРКЕ': 'Объект на проверке',
        'РАЗРАБОТКА СД': 'Разработка СД по объекту',
        'ВКЛ': 'Объект включен в план',
        'ПРЕД. К ИСКЛ': 'Объект предлагается к исключению',
        'ИСКЛ': 'Объект исключен из плана'
    })

    # Создаем сводные таблицы и объединяем
    result = pd.concat([
        pd.crosstab(df['ПО_Общества'], df['План']),
        pd.crosstab(df['ПО_Общества'], df['ДВ']),
        pd.crosstab(df['ПО_Общества'], df['КП']),
        pd.crosstab(df['ПО_Общества'], df['Признак_МТР_в_заказе']),
        pd.crosstab(df['ПО_Общества'], df['Передано_в_ОДСиССР']),
        pd.crosstab(df['ПО_Общества'], df['Направлено_на_осмечивание']),
        pd.crosstab(df['ПО_Общества'], df['Статус_объекта'])
    ], axis=1, sort=False)    

    # Добавляем настоящие столбцы
    required_columns = ['Основной', 'Доп_1', 'Доп_2',
                       'ДВ на проверке', 'ДВ принята в работу', 'ДВ отсутствует',
                       'КП на проверке', 'КП принято в работу', 'КП отсутствует', 'КП не требуется',
                       'Есть признаки МТР в заказе', 'Нет признаков МТР в заказе', 'Не требуется МТР в заказе',
                       'Передано в ОДСиССР', 'Не передано в ОДСиССР',
                       'Направлено на осмечивание', 'Не направлено на осмечивание', 'СД на доработке',
                       'Объект на проверке', 'Разработка СД по объекту', 'Объект включен в план', 'Объект предлагается к исключению', 'Объект исключен из плана']
    for col in required_columns:
        if col not in result.columns:
            result[col] = 0
  
    # Сбрасываем индекс, чтобы 'ПО_Общества' стал обычным столбцом
    result = result.reset_index()

    # Добавляем столбец 'Кол-во объектов' - количество уникальных записей для каждого ПО_Общества
    count_series = df.groupby('ПО_Общества').size().reset_index(name='Кол-во объектов')
    result = result.merge(count_series, on='ПО_Общества', how='left')

    # Формируем итоговую таблицу
    result = result[['ПО_Общества', 'Кол-во объектов'] + required_columns]

    # Добавляем строку с общим итогом
    total_row = result.sum(numeric_only=True)
    total_row['ПО_Общества'] = 'Общий итог'
    result = pd.concat([result, pd.DataFrame([total_row])], ignore_index=True, sort=False)
    
    return result

# [Остальные функции остаются без изменений - create_doughnut_chart_matplotlib, create_status_doughnut_chart, 
# create_status_bar_chart, create_docx_report, set_cell_shading, create_table_with_chart, create_table_without_chart]

def create_doughnut_chart_matplotlib(df, chart_title, sheet_type):
    """Создание кольцевой диаграммы с использованием Matplotlib"""
    # Находим строку с общим итогом
    total_row = df[df['ПО_Общества'] == 'Общий итог']
    if total_row.empty:
        print("Не найдена строка с общим итогом для построения диаграммы")
        return None
    
    # Данные для диаграммы
    labels = ['Основной', 'Доп_1', 'Доп_2']
    sizes = [
        total_row['Основной'].iloc[0],
        total_row['Доп_1'].iloc[0],
        total_row['Доп_2'].iloc[0]
    ]
    
    # Вычисляем общее количество объектов
    total_objects = sum(sizes)
    
    # Цвета для диаграммы
    colors = ['#99ff99', '#66b3ff', '#ff9999']
    
    # Создаем фигуру с увеличенной высотой для размещения легенды под диаграммой
    fig, ax = plt.subplots(figsize=(5.0, 5.5))
    
    # Создаем кольцевую диаграмму
    wedges, texts, autotexts = ax.pie(sizes, labels=None, colors=colors, autopct='%1.1f%%',
                                      startangle=90, radius=1.3, wedgeprops=dict(width=0.7, edgecolor='w', linewidth=2),
                                      pctdistance=0.75)
    
    # Настраиваем внешний вид процентов
    for autotext in autotexts:
        autotext.set_color('#2c3e50')
        autotext.set_fontweight('bold')
        autotext.set_fontsize(12)
    
    # Добавляем центральный текст с общим количеством объектов
    center_text = f'{total_objects}'
    ax.text(0, 0, center_text, ha='center', va='center', fontsize=16, fontweight='bold', color='#2c3e50')
    
    # Добавляем подпись под центральным числом
    ax.text(0, -0.2, 'объектов', ha='center', va='center', fontsize=14, color='#2c3e50')
    
    # Настраиваем заголовок
    ax.set_title(chart_title, fontsize=12, fontweight='bold', pad=15)
    
    # Добавляем легенду ПОД диаграммой
    legend_labels = [f'{label}: {size}' for label, size in zip(labels, sizes)]
    legend = ax.legend(wedges, legend_labels, 
                      loc='upper center', bbox_to_anchor=(0.5, 0.0),
                      ncol=3, fontsize=12)
    
    # Устанавливаем размер шрифта заголовка легенды
    legend.get_title().set_fontsize(11)
    
    # Убеждаемся, что диаграмма круговая
    ax.axis('equal')
    
    # Настраиваем общий вид с учетом легенды
    plt.tight_layout()
    
    # Сохраняем диаграмму в буфер памяти
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', 
                facecolor='#f8f9fa', edgecolor='none')
    buffer.seek(0)
    
    # Закрываем фигуру для освобождения памяти
    plt.close(fig)
    
    return buffer

def create_status_doughnut_chart(labels, sizes, chart_title, colors=None, figsize=(5.0, 5.5)):
    """Создание кольцевой диаграммы для статусов"""
    if colors is None:
        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#c2c2f0', '#ffb3e6', '#c4e17f']
    
    # Вычисляем общее количество
    total = sum(sizes)
    
    # Создаем фигуру с увеличенной высотой для размещения легенды под диаграммой
    fig, ax = plt.subplots(figsize=figsize)
    
    # Создаем кольцевую диаграмму
    wedges, texts, autotexts = ax.pie(sizes, labels=None, colors=colors, autopct='%1.1f%%',
                                      startangle=90, radius=1.3, wedgeprops=dict(width=0.7, edgecolor='w', linewidth=2),
                                      pctdistance=0.75)
    
    # Настраиваем внешний вид процентов
    for autotext in autotexts:
        autotext.set_color('#2c3e50')
        autotext.set_fontweight('bold')
        autotext.set_fontsize(10)
    
    # Добавляем центральный текст с общим количеством
    center_text = f'{total}'
    ax.text(0, 0, center_text, ha='center', va='center', fontsize=16, fontweight='bold', color='#2c3e50')
    
    # Добавляем подпись под центральным числом
    ax.text(0, -0.2, 'объектов', ha='center', va='center', fontsize=14, color='#2c3e50')
    
    # Настраиваем заголовок
    ax.set_title(chart_title, fontsize=12, fontweight='bold', pad=15)
    
    # Добавляем легенду ПОД диаграммой
    legend_labels = [f'{label}: {size}' for label, size in zip(labels, sizes)]
    ncol = 1 if len(labels) <= 4 else 2
    legend = ax.legend(wedges, legend_labels, 
                      loc='upper center', bbox_to_anchor=(0.5, 0.0),
                      ncol=ncol, fontsize=10)
    
    # Убеждаемся, что диаграмма круговая
    ax.axis('equal')
    
    # Настраиваем общий вид с учетом легенды
    plt.tight_layout()
    
    # Сохраняем диаграмму в буфер памяти
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', 
                facecolor='#f8f9fa', edgecolor='none')
    buffer.seek(0)
    
    # Закрываем фигуру для освобождения памяти
    plt.close(fig)
    
    return buffer

def create_status_bar_chart(labels, sizes, chart_title, colors=None, figsize=(10, 6)):
    """Создание горизонтальной столбчатой диаграммы для статусов объектов с сортировкой по убыванию"""
    if colors is None:
        colors = ['#66b3ff', '#99ff99', '#c2c2f0', '#ffcc99', '#ff9999']
    
    # Убедимся, что все значения числовые
    sizes = [float(size) for size in sizes]
    
    # СОРТИРУЕМ ПО ВОЗРАСТАНИЮ (от меньшего к большему) 
    sorted_indices = sorted(range(len(sizes)), key=lambda i: sizes[i])
    
    sorted_labels = [labels[i] for i in sorted_indices]
    sorted_sizes = [sizes[i] for i in sorted_indices]
    sorted_colors = [colors[i] for i in sorted_indices]
    
    # Вычисляем общее количество и проценты для отсортированных данных
    total = sum(sorted_sizes)
    percentages = [f'({size/total*100:.1f}%)' if total > 0 else '(0%)' for size in sorted_sizes]
    
    # Создаем фигуру
    fig, ax = plt.subplots(figsize=figsize)
    
    # Создаем горизонтальную столбчатую диаграмму
    y_pos = range(len(sorted_labels))
    bars = ax.barh(y_pos, sorted_sizes, color=sorted_colors, edgecolor='white', linewidth=1.5, height=0.7)
    
    # Настраиваем внешний вид столбцов
    for i, (bar, size, percentage) in enumerate(zip(bars, sorted_sizes, percentages)):
        width = bar.get_width()
        # Добавляем подписи с количеством и процентом справа от столбца
        ax.text(width + (max(sorted_sizes) * 0.01), bar.get_y() + bar.get_height()/2,
                f'{int(size)} {percentage}', ha='left', va='center', fontsize=11, fontweight='bold')
    
    # Настраиваем заголовок и подписи
    ax.set_title(chart_title, fontsize=14, fontweight='bold', pad=20)
    ax.set_xlabel('Количество объектов', fontsize=12)
    
    # Настраиваем ось Y с метками (отсортированными)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(sorted_labels, fontsize=11)
    
    # Настраиваем внешний вид осей
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#d3d3d3')
    ax.spines['bottom'].set_color('#d3d3d3')
    
    # Настраиваем сетку
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    
    # Добавляем общее количество в заголовок или как подпись
    ax.text(0.98, 0.02, f'Всего объектов: {int(total)}', 
            transform=ax.transAxes, ha='right', va='bottom', 
            fontsize=12, fontweight='bold', bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))
    
    # Настраиваем общий вид
    plt.tight_layout()
    
    # Сохраняем диаграмму в буфер памяти
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', 
                facecolor='#f8f9fa', edgecolor='none')
    buffer.seek(0)
    
    # Закрываем фигуру для освобождения памяти
    plt.close(fig)
    
    return buffer

def create_docx_report(kr_df, totr_df, output_filename=None):
    """Создание полного отчета в формате DOCX со всеми диаграммами и правильным форматированием"""
    
    # Используем выходной файл из конфигурации, если не указан другой
    if output_filename is None:
        output_filename = FILE_PATHS['output_file']
    
    try:
        doc = Document()
        
        # Устанавливаем поля документа
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.05)
            section.bottom_margin = Cm(0.95)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(1.41)
        
        # Настраиваем шрифт Arial для всего документа
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Заголовок отчета
        title = doc.add_paragraph()
        title_run = title.add_run('Отчет по подготовке планов ТОиР на 2027 года')
        title_run.font.size = Pt(16)
        title_run.font.name = 'Arial'
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(6)
        title.paragraph_format.line_spacing = 1
        
        # Дата
        timestamp = datetime.now().strftime("%d.%m.%Y")
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Дата: {timestamp}')
        date_run.font.size = Pt(12)
        date_run.font.name = 'Arial'
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.paragraph_format.space_before = Pt(6)
        date_para.paragraph_format.space_after = Pt(6)
        date_para.paragraph_format.line_spacing = 1
        
        # Раздел КР
        kr_title = doc.add_paragraph()
        kr_title_run = kr_title.add_run('КАПИТАЛЬНЫЙ РЕМОНТ')
        kr_title_run.font.size = Pt(12)
        kr_title_run.font.name = 'Arial'
        kr_title_run.bold = True
        kr_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        kr_title.paragraph_format.space_before = Pt(6)
        kr_title.paragraph_format.space_after = Pt(6)
        kr_title.paragraph_format.line_spacing = 1
        
        # Таблица 1: Количество объектов для КР
        table1_title = doc.add_paragraph()
        table1_title_run = table1_title.add_run('КР: Количество объектов')
        table1_title_run.font.size = Pt(12)
        table1_title_run.font.name = 'Arial'
        table1_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table1_title.paragraph_format.space_before = Pt(6)
        table1_title.paragraph_format.space_after = Pt(0)
        table1_title.paragraph_format.line_spacing = 1
        
        # Создаем таблицу 1 для КР с дополнительным столбцом для диаграммы
        table1_cols = ['ПО_Общества', 'Кол-во объектов', 'Основной', 'Доп_1', 'Доп_2']
        table1_data = kr_df[table1_cols]
        create_table_with_chart(doc, table1_data, "КР: Распределение по планам", "kr_plan_chart", chart_size=(6.06, 6.1))
        
        # Таблица 2.1: Статусы ДВ для КР
        table21_title = doc.add_paragraph()
        table21_title_run = table21_title.add_run('КР: Статусы ДВ')
        table21_title_run.font.size = Pt(12)
        table21_title_run.font.name = 'Arial'
        table21_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table21_title.paragraph_format.space_before = Pt(6)
        table21_title.paragraph_format.space_after = Pt(0)
        table21_title.paragraph_format.line_spacing = 1
        
        table21_cols = ['ПО_Общества', 'ДВ на проверке', 'ДВ принята в работу', 'ДВ отсутствует']
        table21_data = kr_df[table21_cols]
        create_table_with_chart(doc, table21_data, "КР: Статусы ДВ", "kr_dv_chart", chart_size=(5.91, 6.5))
        
        # Таблица 2.2: Статусы КП для КР
        table22_title = doc.add_paragraph()
        table22_title_run = table22_title.add_run('КР: Статусы КП')
        table22_title_run.font.size = Pt(12)
        table22_title_run.font.name = 'Arial'
        table22_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table22_title.paragraph_format.space_before = Pt(6)
        table22_title.paragraph_format.space_after = Pt(0)
        table22_title.paragraph_format.line_spacing = 1
        
        table22_cols = ['ПО_Общества', 'КП на проверке', 'КП принято в работу', 'КП отсутствует', 'КП не требуется']
        table22_data = kr_df[table22_cols]
        create_table_with_chart(doc, table22_data, "КР: Статусы КП", "kr_kp_chart", chart_size=(5.91, 6.5))
        
        # Разрыв страницы
        doc.add_page_break()
        
        # Таблица 2.3: Статусы МТР для КР
        table23_title = doc.add_paragraph()
        table23_title_run = table23_title.add_run('КР: Статусы МТР')
        table23_title_run.font.size = Pt(12)
        table23_title_run.font.name = 'Arial'
        table23_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table23_title.paragraph_format.space_before = Pt(6)
        table23_title.paragraph_format.space_after = Pt(0)
        table23_title.paragraph_format.line_spacing = 1
        
        table23_cols = ['ПО_Общества', 'МТР на проверке', 'ЕСТЬ замечания к МТР', 'Замечаний к МТР НЕТ', 'Внесение МТР не требуется']
        table23_data = kr_df[table23_cols]
        create_table_with_chart(doc, table23_data, "КР: Статусы МТР", "kr_kp_chart", chart_size=(5.91, 6.5))

        # Таблица 2.4: Признаки наличия у заказа ведомости МТР для КР
        table24_title = doc.add_paragraph()
        table24_title_run = table24_title.add_run('КР: Признаки наличия у заказа ведомости МТР')
        table24_title_run.font.size = Pt(12)
        table24_title_run.font.name = 'Arial'
        table24_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table24_title.paragraph_format.space_before = Pt(6)
        table24_title.paragraph_format.space_after = Pt(0)
        table24_title.paragraph_format.line_spacing = 1
        
        table24_cols = ['ПО_Общества', 'Есть признаки МТР в заказе', 'Нет признаков МТР в заказе', 'Не требуется МТР в заказе']
        table24_data = kr_df[table24_cols]
        create_table_with_chart(doc, table24_data, "КР: Признаки наличия у заказа ведомости МТР", "kr_kp_chart", chart_size=(5.91, 6.5))

        # Таблица 3.1: Передача в ОДСиССР для КР
        table31_title = doc.add_paragraph()
        table31_title_run = table31_title.add_run('КР: Передача в ОДСиССР')
        table31_title_run.font.size = Pt(12)
        table31_title_run.font.name = 'Arial'
        table31_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table31_title.paragraph_format.space_before = Pt(6)
        table31_title.paragraph_format.space_after = Pt(0)
        table31_title.paragraph_format.line_spacing = 1
        
        table31_cols = ['ПО_Общества', 'Передано в ОДСиССР', 'Не передано в ОДСиССР']
        table31_data = kr_df[table31_cols]
        create_table_with_chart(doc, table31_data, "КР: Передача в ОДСиССР", "kr_ods_chart", chart_size=(5.91, 6.5))
        
        # Разрыв страницы
        doc.add_page_break()

        # Таблица 3.2: Направление на осмечивание для КР
        table32_title = doc.add_paragraph()
        table32_title_run = table32_title.add_run('КР: Направление на осмечивание')
        table32_title_run.font.size = Pt(12)
        table32_title_run.font.name = 'Arial'
        table32_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table32_title.paragraph_format.space_before = Pt(6)
        table32_title.paragraph_format.space_after = Pt(0)
        table32_title.paragraph_format.line_spacing = 1
        
        table32_cols = ['ПО_Общества', 'Направлено на осмечивание', 'Не направлено на осмечивание', 'СД на доработке']
        table32_data = kr_df[table32_cols]
        create_table_with_chart(doc, table32_data, "КР: Направление на осмечивание", "kr_osmech_chart", chart_size=(5.91, 6.5))
        
        # Таблица 4: Готовность объектов для КР (без диаграммы)
        table4_title = doc.add_paragraph()
        table4_title_run = table4_title.add_run('КР: Готовность объектов')
        table4_title_run.font.size = Pt(12)
        table4_title_run.font.name = 'Arial'
        table4_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table4_title.paragraph_format.space_before = Pt(6)
        table4_title.paragraph_format.space_after = Pt(0)
        table4_title.paragraph_format.line_spacing = 1
        
        table4_cols = ['ПО_Общества', 'Объект на проверке', 'Разработка СД по объекту', 'Объект включен в план', 
                      'Объект предлагается к исключению', 'Объект исключен из плана']
        table4_data = kr_df[table4_cols]
        create_table_without_chart(doc, table4_data)
        
        # Диаграмма Статусы объектов для КР после таблицы 4
        kr_total = kr_df[kr_df['ПО_Общества'] == 'Общий итог'].iloc[0]
        status_labels = ['Объект на проверке', 'Разработка СД по объекту', 'Объект включен в план', 
                        'Объект предлагается к исключению', 'Объект исключен из плана']
        status_data = [kr_total['Объект на проверке'], kr_total['Разработка СД по объекту'], 
                      kr_total['Объект включен в план'], kr_total['Объект предлагается к исключению'], 
                      kr_total['Объект исключен из плана']]
        status_chart_buffer = create_status_bar_chart(status_labels, status_data, "КР: Статусы объектов")
        
        if status_chart_buffer:
            temp_file_path = save_buffer_to_temp_file(status_chart_buffer, "kr_status_chart")
            if temp_file_path and os.path.exists(temp_file_path):
                chart_para = doc.add_paragraph()
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_para.paragraph_format.space_before = Pt(6)
                chart_para.paragraph_format.space_after = Pt(0)
                chart_para.paragraph_format.line_spacing = 1
                run = chart_para.add_run()
                run.add_picture(temp_file_path, width=Cm(15.24), height=Cm(9.02))
        
        # Разрыв страницы для ТОиТР
        doc.add_page_break()
        
        # Раздел ТОиТР
        totr_title = doc.add_paragraph()
        totr_title_run = totr_title.add_run('ТЕХНИЧЕСКОЕ ОБСЛУЖИВАНИЕ И ТЕКУЩИЙ РЕМОНТ')
        totr_title_run.font.size = Pt(12)
        totr_title_run.font.name = 'Arial'
        totr_title_run.bold = True
        totr_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        totr_title.paragraph_format.space_before = Pt(6)
        totr_title.paragraph_format.space_after = Pt(0)
        totr_title.paragraph_format.line_spacing = 1
        
        # Таблица 1 для ТОиТР с дополнительным столбцом для диаграммы
        table1_title_totr = doc.add_paragraph()
        table1_title_run_totr = table1_title_totr.add_run('ТОиТР: Количество объектов')
        table1_title_run_totr.font.size = Pt(12)
        table1_title_run_totr.font.name = 'Arial'
        table1_title_totr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table1_title_totr.paragraph_format.space_before = Pt(6)
        table1_title_totr.paragraph_format.space_after = Pt(0)
        table1_title_totr.paragraph_format.line_spacing = 1
        
        table1_data_totr = totr_df[table1_cols]
        create_table_with_chart(doc, table1_data_totr, "ТОиТР: Распределение по планам", "totr_plan_chart", chart_size=(6.06, 6.1))
        
        # Таблица 2.1: Статусы ДВ для ТОиТР
        table21_title_totr = doc.add_paragraph()
        table21_title_run_totr = table21_title_totr.add_run('ТОиТР: Статусы ДВ')
        table21_title_run_totr.font.size = Pt(12)
        table21_title_run_totr.font.name = 'Arial'
        table21_title_totr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table21_title_totr.paragraph_format.space_before = Pt(6)
        table21_title_totr.paragraph_format.space_after = Pt(0)
        table21_title_totr.paragraph_format.line_spacing = 1
        
        table21_data_totr = totr_df[table21_cols]
        create_table_with_chart(doc, table21_data_totr, "ТОиТР: Статусы ДВ", "totr_dv_chart", chart_size=(5.91, 6.5))
        
        # Таблица 2.2: Статусы КП для ТОиТР
        table22_title_totr = doc.add_paragraph()
        table22_title_run_totr = table22_title_totr.add_run('ТОиТР: Статусы КП')
        table22_title_run_totr.font.size = Pt(12)
        table22_title_run_totr.font.name = 'Arial'
        table22_title_totr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table22_title_totr.paragraph_format.space_before = Pt(6)
        table22_title_totr.paragraph_format.space_after = Pt(0)
        table22_title_totr.paragraph_format.line_spacing = 1
        
        table22_data_totr = totr_df[table22_cols]
        create_table_with_chart(doc, table22_data_totr, "ТОиТР: Статусы КП", "totr_kp_chart", chart_size=(5.91, 6.5))
        
        # Разрыв страницы
        doc.add_page_break()

        # Таблица 2.4: Признаки наличия у заказа ведомости МТР для ТОиТР
        table24_title = doc.add_paragraph()
        table24_title_run = table24_title.add_run('ТОиТР: Признаки наличия у заказа ведомости МТР')
        table24_title_run.font.size = Pt(12)
        table24_title_run.font.name = 'Arial'
        table24_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table24_title.paragraph_format.space_before = Pt(6)
        table24_title.paragraph_format.space_after = Pt(0)
        table24_title.paragraph_format.line_spacing = 1
        
        table24_data_totr = totr_df[table24_cols]
        create_table_with_chart(doc, table24_data_totr, "ТОиТР: Признаки наличия у заказа ведомости МТР", "totr_kp_chart", chart_size=(5.91, 6.5))

        # Таблица 3.1: Передача в ОДСиССР для ТОиТР
        table31_title_totr = doc.add_paragraph()
        table31_title_run_totr = table31_title_totr.add_run('ТОиТР: Передача в ОДСиССР')
        table31_title_run_totr.font.size = Pt(12)
        table31_title_run_totr.font.name = 'Arial'
        table31_title_totr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table31_title_totr.paragraph_format.space_before = Pt(6)
        table31_title_totr.paragraph_format.space_after = Pt(0)
        table31_title_totr.paragraph_format.line_spacing = 1
        
        table31_data_totr = totr_df[table31_cols]
        create_table_with_chart(doc, table31_data_totr, "ТОиТР: Передача в ОДСиССР", "totr_ods_chart", chart_size=(5.91, 6.5))
        
        # Таблица 3.2: Направление на осмечивание для ТОиТР
        table32_title_totr = doc.add_paragraph()
        table32_title_run_totr = table32_title_totr.add_run('ТОиТР: Направление на осмечивание')
        table32_title_run_totr.font.size = Pt(12)
        table32_title_run_totr.font.name = 'Arial'
        table32_title_totr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table32_title_totr.paragraph_format.space_before = Pt(6)
        table32_title_totr.paragraph_format.space_after = Pt(0)
        table32_title_totr.paragraph_format.line_spacing = 1
        
        table32_data_totr = totr_df[table32_cols]
        create_table_with_chart(doc, table32_data_totr, "ТОиТР: Направление на осмечивание", "totr_osmech_chart", chart_size=(5.91, 6.5))
        
        # Разрыв страницы
        doc.add_page_break()

        # Таблица 4: Готовность объектов для ТОиТР (без диаграммы)
        table4_title_totr = doc.add_paragraph()
        table4_title_run_totr = table4_title_totr.add_run('ТОиТР: Готовность объектов')
        table4_title_run_totr.font.size = Pt(12)
        table4_title_run_totr.font.name = 'Arial'
        table4_title_totr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table4_title_totr.paragraph_format.space_before = Pt(6)
        table4_title_totr.paragraph_format.space_after = Pt(0)
        table4_title_totr.paragraph_format.line_spacing = 1
        
        table4_data_totr = totr_df[table4_cols]
        create_table_without_chart(doc, table4_data_totr)
        
        # Диаграмма Статусы объектов для ТОиТР после таблицы 4
        totr_total = totr_df[totr_df['ПО_Общества'] == 'Общий итог'].iloc[0]
        status_labels_totr = ['Объект на проверке', 'Разработка СД по объекту', 'Объект включен в план', 
                             'Объект предлагается к исключению', 'Объект исключен из плана']
        status_data_totr = [totr_total['Объект на проверке'], totr_total['Разработка СД по объекту'], 
                           totr_total['Объект включен в план'], totr_total['Объект предлагается к исключению'], 
                           totr_total['Объект исключен из плана']]
        status_chart_buffer_totr = create_status_bar_chart(status_labels_totr, status_data_totr, "ТОиТР: Статусы объектов")
        
        if status_chart_buffer_totr:
            temp_file_path = save_buffer_to_temp_file(status_chart_buffer_totr, "totr_status_chart")
            if temp_file_path and os.path.exists(temp_file_path):
                chart_para = doc.add_paragraph()
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_para.paragraph_format.space_before = Pt(6)
                chart_para.paragraph_format.space_after = Pt(0)
                chart_para.paragraph_format.line_spacing = 1
                run = chart_para.add_run()
                run.add_picture(temp_file_path, width=Cm(15.24), height=Cm(9.02))
        
        # Сохраняем документ
        doc.save(output_filename)
        print(f"DOCX отчет успешно создан: {output_filename}")
        
    except Exception as e:
        print(f"Ошибка при создании DOCX отчета: {e}")
        import traceback
        traceback.print_exc()

def set_cell_shading(cell, fill_color):
    """Устанавливает заливку ячейки таблицы"""
    try:
        tc_pr = cell._element.tcPr
        if tc_pr is None:
            tc_pr = parse_xml(r'<w:tcPr {}></w:tcPr>'.format(nsdecls('w')))
            cell._element.append(tc_pr)
        
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
        tc_pr.append(shading)
    except Exception as e:
        print(f"Ошибка при установке заливки ячейки: {e}")

def create_table_with_chart(doc, df, chart_title, chart_prefix, chart_size):
    """Создание таблицы с диаграммой"""
    try:
        # Создаем таблицу с дополнительным столбцом для диаграммы
        num_data_cols = len(df.columns)
        num_rows = len(df) + 1  # +1 для заголовка
        table = doc.add_table(rows=num_rows, cols=num_data_cols + 1)  # +1 для столбца с диаграммой
        table.style = 'Table Grid'
        
        # РАСЧЕТ ШИРИН СТОЛБЦОВ
        # Общая доступная ширина листа (A4): 21 см
        # Первый столбец: 3.19 см, последний: 7.46 см
        # Оставшаяся ширина для средних столбцов: 21 - 3.19 - 7.46 = 10.35 см
        first_col_width = Cm(3.19)
        last_col_width = Cm(7.46)
        middle_cols_total_width = Cm(10.35)
        
        # Количество средних столбцов (все столбцы данных кроме первого)
        num_middle_cols = num_data_cols - 1
        
        # Равномерно распределяем оставшуюся ширину между средними столбцами
        if num_middle_cols > 0:
            middle_col_width = middle_cols_total_width / num_middle_cols
        else:
            middle_col_width = Cm(0)
        
        # Формируем список ширин для всех столбцов
        widths = [first_col_width]  # первый столбец
        
        # Добавляем средние столбцы
        for i in range(num_middle_cols):
            widths.append(middle_col_width)
        
        # Добавляем последний столбец для диаграммы
        widths.append(last_col_width)
        
        # Применяем ширины к столбцам таблицы
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Заголовки таблицы (только для данных)
        header_cells = table.rows[0].cells
        column_names = list(df.columns)
        
        for i, column_name in enumerate(column_names):
            # Упрощаем названия столбцов для лучшего отображения
            simplified_name = str(column_name).replace('_', ' ')
            header_cells[i].text = simplified_name
            
            # Форматирование заголовков
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
            
            # Заливка заголовков
            set_cell_shading(header_cells[i], 'F8F9FA')
        
        # Пустой заголовок для столбца с диаграммой
        header_cells[num_data_cols].text = ""  # Пустой заголовок для столбца с диаграммой
        for paragraph in header_cells[num_data_cols].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.name = 'Arial'
        
        set_cell_shading(header_cells[num_data_cols], 'F8F9FA')
        
        # Данные таблицы
        for row_idx, row_data in enumerate(df.itertuples(), 1):
            for col_idx, value in enumerate(row_data[1:], 0):  # Пропускаем индекс
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value)
                
                # Форматирование данных
                for paragraph in cell.paragraphs:
                    if col_idx == 0:  # Первый столбец - выравнивание по левому краю и по центру
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:  # Остальные столбцы - по центру
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.name = 'Arial'
        
        # Форматирование итоговой строки
        total_cells = table.rows[len(df)].cells
        for cell in total_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
            
            # Заливка итоговой строки
            set_cell_shading(cell, 'F8F9FA')
        
        # Объединяем ВСЕ ячейки в последнем столбце для размещения диаграммы
        if num_rows > 1:
            start_cell = table.rows[0].cells[num_data_cols]  # Начинаем с заголовка
            for row_idx in range(1, num_rows):
                next_cell = table.rows[row_idx].cells[num_data_cols]
                start_cell.merge(next_cell)
        
        # Создаем и вставляем соответствующую диаграмму в объединенную ячейку
        total_row = df[df['ПО_Общества'] == 'Общий итог']
        if not total_row.empty:
            if "Распределение по планам" in chart_title:
                # Для таблицы 1 - специальная диаграмма распределения по планам
                chart_buffer = create_doughnut_chart_matplotlib(df, chart_title, "")
            else:
                # Для остальных таблиц - диаграммы статусов
                data_columns = list(df.columns)[1:]  # Исключаем 'ПО_Общества'
                data_values = [total_row[col].iloc[0] for col in data_columns]
                chart_buffer = create_status_doughnut_chart(data_columns, data_values, chart_title)
            
            if chart_buffer:
                temp_file_path = save_buffer_to_temp_file(chart_buffer, chart_prefix)
                if temp_file_path and os.path.exists(temp_file_path):
                    # Вставляем диаграмму в объединенную ячейку
                    cell = table.rows[0].cells[num_data_cols]  # Первая ячейка объединенного столбца
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(temp_file_path, width=Cm(chart_size[0]), height=Cm(chart_size[1]))
                
    except Exception as e:
        print(f"Ошибка при создании таблицы с диаграммой: {e}")

def create_table_without_chart(doc, df):
    """Создание таблицы без диаграммы (для таблицы 4)"""
    try:
        # Создаем таблицу без дополнительного столбца для диаграммы
        num_data_cols = len(df.columns)
        num_rows = len(df) + 1  # +1 для заголовка
        table = doc.add_table(rows=num_rows, cols=num_data_cols)
        table.style = 'Table Grid'
        
        # РАСЧЕТ ШИРИН СТОЛБЦОВ ДЛЯ ТАБЛИЦЫ БЕЗ ДИАГРАММЫ
        # Общая ширина: 21 см (ширина листа A4)
        # Первый столбец: 3.19 см
        # Оставшаяся ширина для остальных столбцов: 21 - 3.19 = 17.81 см
        first_col_width = Cm(3.19)
        remaining_width = Cm(17.81)
        
        # Равномерно распределяем оставшуюся ширину между остальными столбцами
        if num_data_cols > 1:
            other_cols_width = remaining_width / (num_data_cols - 1)
        else:
            other_cols_width = Cm(0)
        
        widths = [first_col_width] + [other_cols_width] * (num_data_cols - 1)
        
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Заголовки таблицы
        header_cells = table.rows[0].cells
        column_names = list(df.columns)
        
        for i, column_name in enumerate(column_names):
            # Упрощаем названия столбцов для лучшего отображения
            simplified_name = str(column_name).replace('_', ' ')
            header_cells[i].text = simplified_name
            
            # Форматирование заголовков
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
            
            # Заливка заголовков
            set_cell_shading(header_cells[i], 'F8F9FA')
        
        # Данные таблицы
        for row_idx, row_data in enumerate(df.itertuples(), 1):
            for col_idx, value in enumerate(row_data[1:], 0):  # Пропускаем индекс
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value)
                
                # Форматирование данных
                for paragraph in cell.paragraphs:
                    if col_idx == 0:  # Первый столбец - выравнивание по левому краю и по центру
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:  # Остальные столбцы - по центру
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.name = 'Arial'
        
        # Форматирование итоговой строки
        total_cells = table.rows[len(df)].cells
        for cell in total_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
            
            # Заливка итоговой строки
            set_cell_shading(cell, 'F8F9FA')
                
    except Exception as e:
        print(f"Ошибка при создании таблицы без диаграммы: {e}")

def create_combined_report():
    """Создание объединенного отчета"""
    try:
        # Выводим информацию о путях для отладки
        print("Текущая рабочая директория:", os.getcwd())
        print("Базовая директория проекта:", BASE_DIR)
        print("Путь к файлу КР:", FILE_PATHS['kr_file'])
        print("Путь к файлу ТОиТР:", FILE_PATHS['totr_file'])
        print("Выходной файл:", FILE_PATHS['output_file'])
        
        # Проверяем существование базовых папок
        if not os.path.exists(os.path.dirname(FILE_PATHS['kr_file'])):
            print(f"Папка КР не найдена: {os.path.dirname(FILE_PATHS['kr_file'])}")
        
        if not os.path.exists(os.path.dirname(FILE_PATHS['totr_file'])):
            print(f"Папка ТОиТР не найдена: {os.path.dirname(FILE_PATHS['totr_file'])}")

        # Генерируем отчеты
        print("Генерация отчета по капитальному ремонту...")
        kr_df = generate_kr_report()
        
        print("Генерация отчета по техническому обслуживанию и текущему ремонту...")
        totr_df = generate_totr_report()
        
        # Создаем новый документ Word
        print("Создание отчета в формате DOCX...")
        create_docx_report(kr_df, totr_df)
        
        print(f"Файл успешно создан: {FILE_PATHS['output_file']}")
        print(f"Обработано строк в КР: {len(kr_df)}")
        print(f"Обработано строк в ТОиТР: {len(totr_df)}")
        
    except FileNotFoundError as e:
        print(f"Ошибка: {e}")
        print("Пожалуйста, проверьте конфигурацию путей в FILE_PATHS")
    except Exception as e:
        print(f"Общая ошибка при создании отчетов: {e}")
        import traceback
        traceback.print_exc()

# Вспомогательная функция для пространств имен XML
def nsdecls(*prefixes):
    return ' '.join(['xmlns:{}="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'.format(prefix) for prefix in prefixes])

# Запускаем создание объединенного отчета
if __name__ == "__main__":
    create_combined_report()

